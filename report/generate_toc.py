from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

out = Path(__file__).resolve().parent
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CONTENTS")
run.bold = True
run.font.size = Pt(16)
run.font.name = "Times New Roman"

doc.add_paragraph()

TAB_POS = Inches(6.0)

# level: 1 = chapter, 2 = subsection, 0 = unnumbered
entries = [
    (1, "Introduction", "2"),
    (1, "Literature Review", "3"),
    (1, "System Model and Problem Formulation", "4"),
    (2, "System Model", "4-5"),
    (2, "Problem Formulation", "6-7"),
    (1, "Proposed Mechanism", "8"),
    (2, "Pseudo Code", "9"),
    (2, "Algorithm Explanation", "10"),
    (2, "Theories and Proofs", "11"),
    (2, "Time Complexity Analysis", "12"),
    (1, "Experimental Results and Graphs", "13"),
    (2, "Simulation Setup", "13-14"),
    (2, "Performance Evaluation", "14-15"),
    (1, "Conclusion", "15"),
    (2, "Conclusion", "15"),
    (0, "References", "15-16"),
    (0, "Individual Contribution", "17-19"),
    (0, "Plagiarism Report", ""),
]

chapter = 0
sub = 0

for level, text, page in entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.tab_stops.add_tab_stop(
        TAB_POS, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )

    if level == 0:
        label = text
        indent = 0
        bold = True
    elif level == 1:
        chapter += 1
        sub = 0
        label = f"{chapter}  {text}"
        indent = 0
        bold = True
    else:
        sub += 1
        label = f"{chapter}.{sub}  {text}"
        indent = 0.35
        bold = False

    p.paragraph_format.left_indent = Inches(indent)

    r = p.add_run(label)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold

    if page:
        p.add_run("\t")
        pr = p.add_run(page)
        pr.font.name = "Times New Roman"
        pr.font.size = Pt(12)

path = out / "Table_of_Contents.docx"
doc.save(path)
print(path)
